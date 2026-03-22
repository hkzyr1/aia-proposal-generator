"""
AIA友邦储蓄险客户方案 — python-docx 文档生成器

将 template.js (Node.js docx 库) 完整转换为 Python 实现。
保持品牌色系、字号、间距等排版规范不变。
"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

import re

from benefit_data import (
    generate_benefit_table,
    simulate_withdrawal,
    format_usd,
    BASE_TOTAL_SURRENDER,
)


def _parse_child_info(needs_text: str):
    """从核心需求文本中解析孩子当前年龄和目标提取年龄。
    例如：'儿子现在3岁，从19岁开始提取' → (3, 19)
    返回 (child_current_age, target_withdraw_age) 或 (None, None)
    """
    child_age = None
    target_age = None

    # 匹配孩子当前年龄: "儿子/女儿/孩子 现在/今年 X岁"
    m = re.search(r'(?:儿子|女儿|孩子|小孩).*?(?:现在|今年|目前).*?(\d+)\s*岁', needs_text)
    if m:
        child_age = int(m.group(1))

    # 匹配目标提取年龄: "从X岁开始提取" 或 "X岁时候开始提取" 或 "X岁开始"
    m2 = re.search(r'(?:从|在).*?(\d+)\s*岁.*?(?:开始|时候|时)?.*?提取', needs_text)
    if m2:
        target_age = int(m2.group(1))

    return child_age, target_age


def _detect_primary_need(needs_lower: str):
    """检测客户核心需求的主要类别。返回字符串标识。
    优先级：教育 > 传承 > 资产隔离 > 储蓄增值 > 退休 > general
    """
    if any(kw in needs_lower for kw in ["教育", "子女", "孩子", "留学", "学费"]):
        return "education"
    if any(kw in needs_lower for kw in ["传承", "遗产", "继承", "家族"]):
        return "inheritance"
    if any(kw in needs_lower for kw in ["资产隔离", "婚前", "离婚", "保全", "债务"]):
        return "asset_isolation"
    if any(kw in needs_lower for kw in ["储蓄", "增值", "理财", "投资"]):
        return "growth"
    if any(kw in needs_lower for kw in ["退休", "养老", "退休金"]):
        return "retirement"
    return "general"

# ============================================================
# 品牌色系常量
# ============================================================
COLORS = {
    "NAVY": "003B73",
    "GOLD": "C8A951",
    "TEXT": "333333",
    "GRAY": "666666",
    "WHITE": "FFFFFF",
    "ROW_ALT_1": "E8F0FE",
    "ROW_ALT_2": "F5F8FC",
    "BORDER": "CCCCCC",
}


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _set_cell_shading(cell, color: str):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_borders(cell, color: str = "CCCCCC", size: int = 4):
    """设置单元格四边边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=80, bottom=80, left=80, right=80):
    """设置单元格内边距 (twips)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _set_cell_width(cell, width_dxa: int):
    """设置单元格宽度"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    tcPr.append(tcW)


def _add_run(paragraph, text, bold=False, italic=False, size=12, color="333333", font="Microsoft YaHei"):
    """添加格式化文本 run"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color)
    run.font.name = font
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    return run


def _create_header_cell(table, row_idx, col_idx, text, width_dxa):
    """创建表头单元格"""
    cell = table.cell(row_idx, col_idx)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, text, bold=True, size=10, color="FFFFFF", font="Microsoft YaHei")
    _set_cell_shading(cell, COLORS["NAVY"])
    _set_cell_borders(cell, COLORS["BORDER"])
    _set_cell_margins(cell)
    _set_cell_width(cell, width_dxa)


def _create_data_cell(table, row_idx, col_idx, text, width_dxa, fill_color, bold=False, align="right"):
    """创建数据单元格"""
    cell = table.cell(row_idx, col_idx)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == "right" else WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, text, bold=bold, size=10, font="Microsoft YaHei")
    _set_cell_shading(cell, fill_color)
    _set_cell_borders(cell, COLORS["BORDER"])
    _set_cell_margins(cell)
    _set_cell_width(cell, width_dxa)


def _add_heading(doc, text, level=1):
    """添加标题段落，使用品牌色"""
    p = doc.add_paragraph()
    _add_run(p, text, bold=True, size=14 if level == 1 else 13, color=COLORS["NAVY"])
    p.paragraph_format.space_before = Pt(12 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 5)
    return p


def _add_body(doc, text, space_after=9, runs=None):
    """添加正文段落，支持多 run"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if runs:
        for r in runs:
            _add_run(p, r.get("text", ""),
                     bold=r.get("bold", False),
                     italic=r.get("italic", False),
                     size=r.get("size", 12),
                     color=r.get("color", COLORS["TEXT"]))
    else:
        _add_run(p, text)
    return p


def _add_numbered_item(doc, text, bold_prefix="", space_after=None):
    """添加编号列表项"""
    p = doc.add_paragraph()
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    if bold_prefix:
        _add_run(p, bold_prefix, bold=True)
        _add_run(p, text)
    else:
        _add_run(p, text)
    # 使用简单编号前缀
    return p


def _add_bullet_item(doc, text, space_after=None):
    """添加项目符号列表项"""
    p = doc.add_paragraph()
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    _add_run(p, f"• {text}")
    return p


def _add_page_break(doc):
    """添加分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    br_elem = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(br_elem)


def generate_proposal(
    client_name: str,
    client_age: int,
    client_gender: str,
    client_occupation: str,
    client_income: str,
    client_family: str,
    client_needs: str,
    annual_premium: int,
    retirement_age: int = None,
    custom_notes: str = "",
) -> io.BytesIO:
    """
    生成完整的客户方案 .docx 文档，返回 BytesIO 对象。

    Parameters:
        client_name: 客户姓名
        client_age: 客户年龄
        client_gender: 性别
        client_occupation: 职业
        client_income: 年收入描述
        client_family: 家庭结构
        client_needs: 核心需求
        annual_premium: 年缴保费（美元整数）
        retirement_age: 目标退休年龄（可选）
        custom_notes: 自定义备注
    """
    doc = Document()

    # ---- 全局默认样式 ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    font.color.rgb = _rgb(COLORS["TEXT"])
    # 设置东亚字体
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # 计算常用值
    total_premium = annual_premium * 5
    scale = annual_premium / 10000
    proposal_date = datetime.now().strftime("%Y年%m月%d日")
    title_suffix = "先生" if client_gender == "男" else "女士"

    # 退休相关
    if retirement_age and retirement_age > client_age:
        retirement_year = retirement_age - client_age  # 保单年度
    else:
        retirement_year = None

    # 检测核心需求类别和子女信息
    needs_lower_early = client_needs.lower() if client_needs else ""
    primary_need = _detect_primary_need(needs_lower_early)
    child_current_age, child_target_age = _parse_child_info(client_needs or "")

    # 教育金提取年份（保单年度）
    education_withdrawal_year = None
    if primary_need == "education" and child_current_age is not None and child_target_age is not None:
        education_withdrawal_year = child_target_age - child_current_age
        if education_withdrawal_year < 1:
            education_withdrawal_year = None

    # ============================================================
    # 模块 1：封面页
    # ============================================================
    # 空行
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(30)

    # 客户姓名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, client_name, bold=True, size=26, color=COLORS["NAVY"])

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    _add_run(p, "个人财富管理方案", bold=True, size=20, color=COLORS["NAVY"])

    # 金色分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p_element = p._element
    pPr = p_element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{COLORS["GOLD"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    _add_run(p, " ", size=12)

    # 日期
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    _add_run(p, f"方案日期：{proposal_date}", size=11)

    # AIA
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    _add_run(p, "AIA 友邦保险集团", bold=True, size=11, color=COLORS["NAVY"])

    # ============================================================
    # 分页进入正文 Section
    # ============================================================
    _add_page_break(doc)

    # ---- 添加页眉 ----
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("AIA 友邦保险")
    hr.font.size = Pt(9)
    hr.font.color.rgb = _rgb(COLORS["GRAY"])
    hr.font.name = "Arial"

    # ============================================================
    # 模块 2：友邦集团介绍（固定内容）
    # ============================================================
    _add_heading(doc, "您的财富守护者 — AIA友邦保险")

    _add_body(doc, "", runs=[
        {"text": "把积蓄交给一家保险公司，最重要的是什么？是这家公司经得起时间的考验。友邦保险"},
        {"text": "创立于1919年", "bold": True},
        {"text": "，跨越了两次世界大战、多次金融危机，至今已稳健运营超过"},
        {"text": "100年", "bold": True},
        {"text": "。它不仅是亚太地区最大的独立上市人寿保险集团，更是全球保险业的标杆。"},
    ])

    _add_body(doc, "", runs=[
        {"text": "顶级信用评级，三大机构一致认可：", "bold": True},
        {"text": "标准普尔 "},
        {"text": "AA（稳定）", "bold": True},
        {"text": "、穆迪 "},
        {"text": "Aa2（稳定）", "bold": True},
        {"text": "、惠誉 "},
        {"text": "AA（稳定）", "bold": True},
        {"text": "。在全球保险公司中，能同时获得三大评级机构AA级认可的屈指可数，友邦便是其中之一。"},
    ])

    _add_body(doc, "", runs=[
        {"text": "强大的资产管理实力：", "bold": True},
        {"text": "集团总投资规模达"},
        {"text": "3,092亿美元", "bold": True},
        {"text": "，投资组合以固定收益类资产为主导——政府债券、优质企业债券占比接近"},
        {"text": "70%", "bold": True},
        {"text": "，辅以全球股票、投资物业等多元配置。这种「稳健为主、增长为辅」的投资策略，正是友邦保单长期分红稳定的底层保障。"},
    ])

    _add_body(doc, "", runs=[
        {"text": "持续增长的业务表现：", "bold": True},
        {"text": "2025年第三季度新业务价值上升25%至"},
        {"text": "14.76亿美元，创历史新高", "bold": True},
        {"text": "；新业务利润率达58.2%。集团资本覆盖率"},
        {"text": "254%", "bold": True},
        {"text": "，远超监管要求，偿付能力极为充裕。在香港，友邦拥有超过"},
        {"text": "360万客户", "bold": True},
        {"text": "，MDRT百万圆桌会员人数连续"},
        {"text": "23年", "bold": True},
        {"text": "全港澳排名第一。"},
    ])

    _add_page_break(doc)

    # ============================================================
    # 模块 3：客户背景摘要
    # ============================================================
    _add_heading(doc, "客户背景摘要")

    # 构建背景描述
    bg_parts = []
    bg_parts.append(f"{client_name}，{client_age}岁")
    if client_gender:
        bg_parts[0] += f"，{client_gender}性"
    if client_family:
        bg_parts.append(f"家庭情况：{client_family}")
    if client_occupation:
        bg_parts.append(f"职业：{client_occupation}")
    if client_income:
        bg_parts.append(f"年收入：{client_income}")
    if client_needs:
        bg_parts.append(f"核心需求：{client_needs}")

    bg_text = "。".join(bg_parts) + "。"
    _add_body(doc, bg_text, space_after=12)

    # ============================================================
    # 模块 4：需求分析（根据客户核心需求动态生成）
    # ============================================================
    _add_heading(doc, "需求分析")

    needs_lower = client_needs.lower() if client_needs else ""

    # --- 动态开头段落（优先按核心需求关键词，退休作为兜底） ---
    if primary_need == "education":
        child_info = f"作为{client_family}的家长，" if client_family else ""
        if child_current_age is not None and child_target_age is not None:
            years_until = child_target_age - child_current_age
            _add_body(doc, "", runs=[
                {"text": f"{child_info}{client_name}深知教育投资对下一代的重要性。孩子目前{child_current_age}岁，距离{child_target_age}岁的关键教育节点还有"},
                {"text": f"整整{years_until}年", "bold": True},
                {"text": "。在优质教育资源日益稀缺、国际学校和海外留学费用持续攀升的今天，提前做好教育金规划，是给孩子最有远见的礼物。"},
            ])
        else:
            _add_body(doc, f"{child_info}{client_name}深知教育投资对下一代的重要性。在优质教育资源日益稀缺、国际学校和海外留学费用持续攀升的今天，提前做好教育金规划，是给孩子最有远见的礼物。")
    elif primary_need == "inheritance":
        _add_body(doc, f"财富的意义不仅在于当下的享有，更在于跨代的传递。{client_name}正处于事业与家庭的成熟期，此时开始规划财富传承，既有充裕的时间让资产增值，也能从容安排分配方案，确保财富按照您的心意惠及后代。")
    elif primary_need == "asset_isolation":
        _add_body(doc, f"在当前复杂的经济与法律环境下，个人资产的安全性不容忽视。{client_name}希望通过合理的金融工具实现资产隔离与保全，这是极具前瞻性的财务决策。香港保险作为境外合法金融资产，在资产保全方面具有独特的制度优势。")
    elif primary_need == "growth":
        _add_body(doc, f"在全球低利率与通胀并存的环境下，{client_name}正在寻找一种兼具安全性与增长性的储蓄方式。传统银行存款收益持续走低，而股市波动又让人难以安心。如何让辛苦积累的财富稳健增值，是当下最迫切的命题。")
    elif primary_need == "retirement" and retirement_year:
        _add_body(doc, "", runs=[
            {"text": f"{client_age}岁到{retirement_age}岁，这是{client_name}规划退休的"},
            {"text": f"黄金窗口期——整整{retirement_year}年", "bold": True},
            {"text": "。在这段时间里，收入处于职业生涯的高峰，储蓄能力最强；而足够的时间也能让一笔合理的投入通过复利实现显著增值。"},
        ])
    elif retirement_year:
        _add_body(doc, "", runs=[
            {"text": f"{client_age}岁到{retirement_age}岁，这是{client_name}规划退休的"},
            {"text": f"黄金窗口期——整整{retirement_year}年", "bold": True},
            {"text": "。在这段时间里，收入处于职业生涯的高峰，储蓄能力最强；而足够的时间也能让一笔合理的投入通过复利实现显著增值。"},
        ])
    else:
        _add_body(doc, f"基于对{client_name}个人情况的深入了解，我们对您当前的财务状况和未来规划需求进行了全面分析，以确保推荐的方案真正契合您的期望与目标。")

    _add_body(doc, "", runs=[
        {"text": "我们为您识别了以下核心规划要点：", "bold": True},
    ])

    # --- 根据客户核心需求动态匹配要点 ---
    # 预定义所有可能的需求要点
    all_need_items = {
        "retirement": ("退休收入保障：", f"按照当前生活水准，退休后需要持续稳定的现金流维持品质生活。社保养老金覆盖范围有限，尤其对于高收入人群而言，仅靠社保远不足以维持退休前的生活品质，需要通过个人储蓄和投资来填补缺口。"),
        "education": ("子女教育金规划：", "优质教育是给孩子最好的投资，但国际学校、海外留学的费用逐年递增。提前以美元配置教育储备金，既能对冲人民币贬值风险，又能通过长期复利让教育基金在需要时已充分增值，从容应对未来的学费支出。"),
        "inheritance": ("财富传承规划：", "财富传承不仅是金额的传递，更是对家人的责任安排。通过保险工具实现跨代财富转移，可以灵活设定受益人、借助保单分拆功能按需分配，且在法律层面具有明确的权属保障，避免未来可能的继承纠纷。"),
        "asset_isolation": ("资产隔离与保全：", "在商业经营或婚姻关系中，个人资产可能面临潜在的法律风险。香港保单作为境外合法金融资产，在资产保全方面具有独特的制度安排，能够有效实现个人财富与经营风险、婚姻风险的合理隔离。"),
        "currency": ("资产币种多元化：", "收入和积蓄以人民币计价为主，缺乏外币资产配置。在全球经济不确定性增大、汇率波动加剧的背景下，配置美元资产是分散单一货币风险的专业选择，也是全球高净值人群的标准做法。"),
        "growth": ("资产稳健增值：", "在银行存款利率持续下行的环境中，传统储蓄方式的实际购买力正被通胀悄然侵蚀。选择一种既有保底保障又能分享长期投资收益的工具，是让财富跑赢通胀、实现稳健增值的关键。"),
        "safety": ("资产安全与增值平衡：", "财务规划需要兼顾安全性与增长性——既不能过于激进承受不必要的风险，也不能过于保守让通胀侵蚀购买力。友邦以七成固定收益加三成增长型资产的配置策略，在安全与收益间取得了最佳平衡。"),
    }

    # 根据 client_needs 关键词匹配
    need_items = []
    if any(kw in needs_lower for kw in ["退休", "养老", "退休金"]):
        need_items.append(all_need_items["retirement"])
    if any(kw in needs_lower for kw in ["教育", "子女", "孩子", "留学", "学费"]):
        need_items.append(all_need_items["education"])
    if any(kw in needs_lower for kw in ["传承", "遗产", "继承", "家族", "下一代"]):
        need_items.append(all_need_items["inheritance"])
    if any(kw in needs_lower for kw in ["资产隔离", "婚前", "离婚", "保全", "债务", "隔离"]):
        need_items.append(all_need_items["asset_isolation"])
    if any(kw in needs_lower for kw in ["货币", "汇率", "美元", "外币", "人民币"]):
        need_items.append(all_need_items["currency"])
    if any(kw in needs_lower for kw in ["储蓄", "增值", "理财", "投资", "收益"]):
        need_items.append(all_need_items["growth"])

    # 如果有退休年龄但没有匹配到退休需求，自动添加（但不在教育金为主需求时强制排第一）
    if retirement_year and not any(item == all_need_items["retirement"] for item in need_items):
        if primary_need == "education":
            need_items.append(all_need_items["retirement"])  # 教育为主时，退休放后面
        else:
            need_items.insert(0, all_need_items["retirement"])

    # 所有场景都加「资产安全与增值平衡」作为兜底，除非已有3个以上
    if len(need_items) < 3 and all_need_items["safety"] not in need_items:
        need_items.append(all_need_items["safety"])

    # 如果需求要点少于2个，补充「资产币种多元化」
    if len(need_items) < 2 and all_need_items["currency"] not in need_items:
        need_items.insert(0, all_need_items["currency"])

    # 兜底：如果完全没有匹配，给出默认3条
    if not need_items:
        need_items = [
            all_need_items["currency"],
            all_need_items["growth"],
            all_need_items["safety"],
        ]

    for i, (prefix, body) in enumerate(need_items):
        sa = 12 if i == len(need_items) - 1 else None
        _add_numbered_item(doc, body, bold_prefix=f"{i+1}. {prefix}", space_after=sa)

    # --- 动态衔接语 ---
    matched_advantages = []
    if any(kw in needs_lower for kw in ["退休", "养老"]):
        matched_advantages.append("灵活提取匹配退休节奏")
    if any(kw in needs_lower for kw in ["教育", "子女", "留学"]):
        matched_advantages.append("灵活提取覆盖各阶段教育支出")
    if any(kw in needs_lower for kw in ["传承", "遗产", "家族"]):
        matched_advantages.append("保单分拆和传承功能实现跨代财富转移")
    if any(kw in needs_lower for kw in ["资产隔离", "保全"]):
        matched_advantages.append("境外保单提供资产保全制度优势")
    # 默认优势
    base_advantages = "美元计价对冲货币风险、长期复利稳健增值"
    if matched_advantages:
        full_advantages = base_advantages + "、" + "、".join(matched_advantages)
    else:
        full_advantages = base_advantages + "、灵活提取匹配人生节奏、保单分拆和传承功能为未来预留充分弹性"

    _add_body(doc, f"基于以上分析，我们为您推荐友邦保险最新一代储蓄计划——环宇盈活储蓄保险计划。它的设计理念与您的需求高度吻合：{full_advantages}。", space_after=12)

    _add_page_break(doc)

    # ============================================================
    # 模块 5：方案推荐概览
    # ============================================================
    _add_heading(doc, "方案推荐概览")

    product_name = "环宇盈活储蓄保险计划（GlobalFlexi Savings Insurance Plan）"
    _add_body(doc, f"推荐产品：{product_name}", runs=[
        {"text": f"推荐产品：{product_name}", "bold": True},
    ], space_after=6)
    _add_body(doc, "", runs=[{"text": "缴费年期：5年", "bold": True}], space_after=6)
    _add_body(doc, "", runs=[{"text": f"年缴保费：{format_usd(annual_premium)}美元", "bold": True}], space_after=6)
    _add_body(doc, "", runs=[{"text": f"总保费：{format_usd(total_premium)}美元", "bold": True}], space_after=6)
    _add_body(doc, "", runs=[{"text": "保单货币：美元", "bold": True}], space_after=6)

    # 概要描述
    year10_total = round(BASE_TOTAL_SURRENDER[10] * scale)
    year10_ratio = round(year10_total / total_premium, 2)

    overview_text = (
        f"每年投入{format_usd(annual_premium)}美元，连续5年，总计{format_usd(total_premium)}美元。"
        f"完成缴费后无需再投入任何资金，保单将在友邦专业投资管理下持续增值。"
    )
    if primary_need == "education" and education_withdrawal_year and education_withdrawal_year in BASE_TOTAL_SURRENDER:
        edu_total = round(BASE_TOTAL_SURRENDER[education_withdrawal_year] * scale)
        edu_ratio = round(edu_total / total_premium, 2)
        overview_text += (
            f"到孩子{child_target_age}岁时（第{education_withdrawal_year}年），预期总价值约{format_usd(edu_total)}美元，"
            f"约为初始投入的{edu_ratio}倍，届时可开始提取教育金。"
        )
    elif retirement_year:
        ret_total = round(BASE_TOTAL_SURRENDER[retirement_year] * scale)
        ret_ratio = round(ret_total / total_premium, 2)
        withdrawal_amt = round(ret_total * 0.065)
        overview_text += (
            f"到您{retirement_age}岁时（第{retirement_year}年），预期总价值约{format_usd(ret_total)}美元，"
            f"约为初始投入的{ret_ratio}倍。"
        )
    _add_body(doc, overview_text, space_after=15)

    _add_page_break(doc)

    # ============================================================
    # 模块 6：收益演示表格
    # ============================================================
    _add_heading(doc, "收益演示")

    _add_body(doc, "单位：美元（基于预期投资回报率演示，非保证）", runs=[
        {"text": "单位：美元（基于预期投资回报率演示，非保证）", "italic": True, "size": 10},
    ], space_after=6)

    # 确定关键年份
    key_years = [10, 20, 30, 40, 50, 100]
    if primary_need == "education" and education_withdrawal_year and education_withdrawal_year not in key_years:
        key_years.append(education_withdrawal_year)
        key_years.sort()
    if retirement_year and retirement_year not in key_years:
        key_years.append(retirement_year)
        key_years.sort()

    benefit_data = generate_benefit_table(annual_premium, key_years)

    # 创建表格
    n_rows = len(benefit_data) + 1
    table = doc.add_table(rows=n_rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_w = 1805
    headers = ["保单年度", "已缴总保费", "保证现金价值", "非保证红利", "预期总价值"]
    for i, h in enumerate(headers):
        _create_header_cell(table, 0, i, h, col_w)

    for row_idx, row_data in enumerate(benefit_data):
        fill = COLORS["ROW_ALT_1"] if row_idx % 2 == 0 else COLORS["ROW_ALT_2"]
        is_bold = row_idx % 2 == 0
        year_label = f"第{row_data['year']}年"
        if primary_need == "education" and education_withdrawal_year and row_data['year'] == education_withdrawal_year:
            year_label += f"\n(孩子{child_target_age}岁)"
        elif retirement_year and row_data['year'] == retirement_year:
            year_label += f"\n({retirement_age}岁)"

        _create_data_cell(table, row_idx + 1, 0, year_label, col_w, fill, bold=is_bold, align="center")
        _create_data_cell(table, row_idx + 1, 1, format_usd(row_data['premium_paid']), col_w, fill)
        _create_data_cell(table, row_idx + 1, 2, format_usd(row_data['guaranteed_cv']), col_w, fill)
        _create_data_cell(table, row_idx + 1, 3, format_usd(row_data['non_guaranteed_bonus']), col_w, fill)
        _create_data_cell(table, row_idx + 1, 4, format_usd(row_data['total_value']), col_w, fill)

    # 收益点评
    year50_total = round(BASE_TOTAL_SURRENDER[50] * scale)
    year100_total = round(BASE_TOTAL_SURRENDER[100] * scale)
    year50_ratio = round(year50_total / total_premium, 2)
    year100_ratio = round(year100_total / total_premium, 2)

    _add_body(doc, "", runs=[
        {"text": f"以上数据清晰地展示了时间与复利的力量。您的{format_usd(total_premium)}美元总投入，到第50年预期增长至约"},
        {"text": f"{format_usd(year50_total)}美元（{year50_ratio}倍）", "bold": True},
        {"text": "；若保单持续至第100年，预期总价值更将增长至约"},
        {"text": f"{format_usd(year100_total)}美元（{year100_ratio}倍）", "bold": True},
        {"text": "——这就是时间与复利的惊人力量。"},
    ], space_after=12)

    _add_page_break(doc)

    # ============================================================
    # 模块 7：提取演示（根据核心需求选择教育金/退休提取）
    # ============================================================
    # 确定提取起始年份和场景
    withdrawal_start_year = None
    withdrawal_scenario = None  # "education" or "retirement"

    if primary_need == "education" and education_withdrawal_year:
        withdrawal_start_year = education_withdrawal_year
        withdrawal_scenario = "education"
    elif retirement_year:
        withdrawal_start_year = retirement_year
        withdrawal_scenario = "retirement"

    if withdrawal_start_year:
        _add_heading(doc, "提取演示")

        withdrawal_result = simulate_withdrawal(annual_premium, withdrawal_start_year)
        annual_w = withdrawal_result["annual_withdrawal"]

        # 根据核心需求 + 提取年龄动态生成描述
        if withdrawal_scenario == "education":
            if child_current_age is not None and child_target_age is not None:
                goal_text = f"为孩子储备教育金、并在孩子{child_target_age}岁时（您{client_age + withdrawal_start_year}岁）开始提取"
            else:
                goal_text = f"为子女储备教育金、并在您{client_age + withdrawal_start_year}岁时开始提取"
        elif primary_need == "inheritance":
            goal_text = f"实现财富传承规划、并在您{client_age + withdrawal_start_year}岁时开始提取"
        elif primary_need == "asset_isolation":
            goal_text = f"实现资产保全与增值、并在您{client_age + withdrawal_start_year}岁时开始提取"
        elif primary_need == "growth":
            goal_text = f"实现资产稳健增值、并在您{client_age + withdrawal_start_year}岁时开始提取"
        else:
            goal_text = f"在您{client_age + withdrawal_start_year}岁时开始提取"

        _add_body(doc, "", runs=[
            {"text": f"基于您{goal_text}的目标，我们按照"},
            {"text": f"第{withdrawal_start_year}年预期总价值 × 6.5%", "bold": True},
            {"text": f"的方式确定年提取金额，从第{withdrawal_start_year}年起每年固定提取"},
            {"text": f"${format_usd(annual_w)}美元", "bold": True},
            {"text": "，以下是提取后保单剩余价值的模拟："},
        ])

        _add_body(doc, "单位：美元（基于预期投资回报率演示，非保证）", runs=[
            {"text": "单位：美元（基于预期投资回报率演示，非保证）", "italic": True, "size": 10},
        ], space_after=6)

        # 选择展示的关键年份
        all_proj = withdrawal_result["projections"]
        show_years_set = set()
        show_years_set.add(withdrawal_start_year)  # 起始年
        # 每隔5年/10年选一些
        for p in all_proj:
            y = p["year"]
            if y % 5 == 0 or y == 100:
                show_years_set.add(y)
        # 确保包含100
        show_years_set.add(100)
        show_years = sorted(show_years_set)

        # 限制展示行数 (最多 ~12行)
        if len(show_years) > 12:
            essential = [withdrawal_start_year, 100]
            remaining = [y for y in show_years if y not in essential]
            step = max(1, len(remaining) // 10)
            selected = remaining[::step]
            show_years = sorted(set(essential + selected))

        proj_dict = {p["year"]: p["balance"] for p in all_proj}

        # 提取表格
        w_table = doc.add_table(rows=len(show_years) + 1, cols=4)
        w_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        w_col = 2256
        w_headers = ["保单年度", f"{client_name}年龄", "年提取金额", "提取后剩余价值"]
        for i, h in enumerate(w_headers):
            _create_header_cell(w_table, 0, i, h, w_col)

        for row_idx, y in enumerate(show_years):
            fill = COLORS["ROW_ALT_1"] if row_idx % 2 == 0 else COLORS["ROW_ALT_2"]
            age_at_year = client_age + y
            balance = proj_dict.get(y, 0)

            _create_data_cell(w_table, row_idx + 1, 0, f"第{y}年", w_col, fill, align="center")
            _create_data_cell(w_table, row_idx + 1, 1, f"{age_at_year}岁", w_col, fill, align="center")
            _create_data_cell(w_table, row_idx + 1, 2, format_usd(annual_w), w_col, fill)
            _create_data_cell(w_table, row_idx + 1, 3, format_usd(balance), w_col, fill, bold=(row_idx == 0))

        # 计算累计提取
        total_withdrawn = annual_w * (100 - withdrawal_start_year + 1)
        year100_balance = proj_dict.get(100, 0)
        total_withdrawn_ratio = round(total_withdrawn / total_premium, 1)

        withdraw_age_label = client_age + withdrawal_start_year
        if withdrawal_scenario == "education" and child_target_age is not None:
            start_desc = f"从孩子{child_target_age}岁（您{withdraw_age_label}岁）起"
            tail_text = "，教育金提取完毕后，保单仍可继续作为退休储备或财富传承工具。"
        else:
            start_desc = f"从您{withdraw_age_label}岁起"
            tail_text = "，真正实现了「取之不尽」的终身现金流。"

        _add_body(doc, "", runs=[
            {"text": f"模拟数据显示，{start_desc}每年固定提取${format_usd(annual_w)}美元，"},
            {"text": "保单剩余价值不仅不会缩水，反而持续增长", "bold": True},
            {"text": f"。到第100年，剩余价值仍高达"},
            {"text": f"${format_usd(year100_balance)}美元", "bold": True},
            {"text": f"。期间累计提取总额达"},
            {"text": f"${format_usd(total_withdrawn)}美元（总投入的{total_withdrawn_ratio}倍）", "bold": True},
            {"text": tail_text},
        ], space_after=9)

        _add_body(doc, "", runs=[
            {"text": "灵活调节的空间：", "bold": True},
            {"text": "以上仅为固定提取的演示场景。实际操作中，您完全可以根据需要灵活调整——某些年份少提或不提，让保单继续积累；需要较大支出时一次性提取更多。"},
        ], space_after=12)

        _add_page_break(doc)

    # ============================================================
    # 模块 8：方案亮点与优势
    # ============================================================
    _add_heading(doc, "方案亮点与优势")

    advantages = [
        ("美元资产配置，构建全球化财富防线",
         "目前的收入和积蓄以人民币为主，这在全球经济波动加剧的今天意味着单一货币风险。这份美元保单帮您在人民币资产之外建立一道坚实的外币屏障。无论未来汇率如何变化，您都拥有一笔以全球储备货币计价的核心资产。"),
        ("底层资产稳健，分红有坚实基础",
         "友邦的分红保单背后是一个规模达3,092亿美元的投资组合。其中固定收益类资产占比接近70%，确保稳健回报；同时配置全球股票、投资物业等资产以捕获增长机遇。这种「七成稳健 + 三成增长」的资产配置结构，正是保单能够长期稳定分红的根本保障。"),
        ("灵活提取，完美匹配退休节奏",
         "保单在缴费期满后即可灵活提取。金额和频率完全由您决定——不需要的年份可暂停提取让资金继续增值，需要大额支出时也可一次性提取更多。这种灵活性是银行定期存款和年金产品无法比拟的。"),
        ("保单分拆与传承，为未来预留无限可能",
         "环宇盈活支持保单分拆功能——未来您可以将一份保单拆分为多份，灵活安排不同用途。同时，保单可指定受益人，实现财富的定向传承。"),
        ("友邦百年品牌，偿付能力远超行业标准",
         "友邦保险以超过100年的历史、三大评级机构一致AA级认可、254%的资本覆盖率（远超监管要求），以及360万香港客户的共同选择，给您最坚实的信心保障。"),
    ]

    for title, body in advantages:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, title, bold=True, size=13, color=COLORS["NAVY"])

        _add_body(doc, body, space_after=10)

    _add_page_break(doc)

    # ============================================================
    # 模块 9：投保流程与注意事项（固定内容）
    # ============================================================
    _add_heading(doc, "投保流程与注意事项")

    _add_body(doc, "一切准备就绪，接下来只需简单几步，即可开启您的财富管理新篇章。")

    # 赴港投保流程
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    _add_run(p, "赴港投保流程", bold=True, size=13, color=COLORS["NAVY"])

    steps = [
        "1. 预约赴港投保时间，我们将为您安排专属接待",
        "2. 携带所需文件前往友邦服务中心（尖沙咀海港城/中环友邦中心）",
        "3. 签署投保文件并缴纳首期保费",
        "4. 保单审批通过后正式生效，通常3-5个工作日",
    ]
    for s in steps:
        _add_body(doc, s, space_after=3)

    # 所需文件
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    _add_run(p, "所需文件", bold=True, size=13, color=COLORS["NAVY"])

    docs_needed = ["港澳通行证（有效期内）", "中国居民身份证", "入境小票（过关时请务必保留）"]
    for d in docs_needed:
        _add_bullet_item(doc, d, space_after=3)

    # 缴费方式
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    _add_run(p, "缴费方式", bold=True, size=13, color=COLORS["NAVY"])

    _add_body(doc, "支持现金、信用卡（Visa/Mastercard）、银行电汇、香港银行本票、香港银行支票等多种缴费方式，灵活便捷。首期保费可在签约当日以信用卡支付，后续保费可通过银行转账安排自动缴付。", space_after=10)

    # 后续服务
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    _add_run(p, "后续服务", bold=True, size=13, color=COLORS["NAVY"])

    _add_body(doc, "投保完成后，您将享有专属客户经理全程服务，包括保单年度检视、保单分拆与货币转换协助、提取与理赔支持、受益人变更等全方位管理。无论您身在内地还是海外，我们都确保您的保单服务畅通无阻。", space_after=10)

    # ============================================================
    # 输出
    # ============================================================
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
